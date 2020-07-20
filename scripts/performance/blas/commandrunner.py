#!/usr/bin/env python3
'''
This tool is meant to manage running commands related to a single project,
but from multiple versions of that project and/or multiple hardware configurations.

Therefore, it is assumed that there is a fixed set of command line arguments that
define two runs that are considered equivalent (e.g. array sizes), and a second set
of command line arguments that change based on hardware/project (e.g. output location).

This tool is also meant to run equivalent commands multiple times in order
to statistically test the validity of the results. Each command must be able to
execute on its own, in any order with respect to other commands.

The tool can be run in multiple stages.
    1) Run the executable and write the results to disk
    2) Statistically analyze/process results on disk
    3) Generate plots
    4) Optionally Interact with the plots/data
    5) Write a summary document
Each stage depends on the previous ones, but they can be executed independently
to aid in development of the final output document. It is also possible to run
the executable on multiple machines if they are all connected to the same network
drive, or the output folders are collected onto a single machine.

To use this tool, first write a class that inherits from ArgumentSetABC that is
specific to the executable being benchmarked. Define all of the arguments, as
well as whether or not they are required/have default values. Second, use the
standard argparse module to define any project specific user inputs, and then pass
the parser into parse_input_arguments(parser). Create an instance of CommandRunner
using the arguments and add ArgumentSets and/or Comparisons.

Optionally write classes that inherit from RunConfiguration and Comparison.
Use a custom RunConfiguration to add project specific arguments that change
the runtime environment of the executable. Use a custom Comparison class to
create more meaningful plots/figures.

This file is setup to act as an example benchmark tool for the Linux commandline
tool "df". See the end of the file for example code.
'''
import argparse
from collections import OrderedDict
import copy
import datetime
import getpass
import hashlib
import itertools
import json
import os
import random
import shutil
import string
import subprocess
import sys
import time

import getspecs

# Optional modules
try:
    import matplotlib.pyplot as plt
    from matplotlib.ticker import FuncFormatter
except ImportError:
    plt = None
try:
    import numpy as np
except ImportError:
    np = None
try:
    import pandas as pd
    if plt is not None:
        pd.plotting.register_matplotlib_converters()
except ImportError:
    pd = None
try:
    import pylatex
except ImportError:
    pylatex = None
try:
    import docx
except ImportError:
    docx = None
try:
    from io import BytesIO
except ImportError:
    BytesIO = None

# Install dependent modules
smi = None
smi_imported = False
def import_rocm_smi(install_path):
    global smi
    global smi_imported
    if not smi_imported:
        smi_imported = True
        try:
            sys.path.append(os.path.join(install_path, 'bin'))
            import rocm_smi
            smi = rocm_smi
        except ImportError:
            print('WARNING - rocm_smi.py not found!')
    return smi

class SystemMonitor(object):
    supported_metrics = [
            'used_memory_percent',
            'fclk_megahertz',
            'mclk_megahertz',
            'sclk_megahertz',
            'socclk_megahertz',
            # 'dcefclk_megahertz',
            'fan_speed_percent',
            ]
    def __init__(self, metrics = supported_metrics):
        if not smi_imported:
            raise RuntimeError('import_rocm_smi(install_path) must be called before consturcting a SystemMonitor')
        if len(metrics) == 0:
            raise ValueError('SystemMonitor must record at least one metric')
        self.metrics = metrics
        self.data = {metric:{} for metric in self.metrics}

    def record_line(self):
        now = datetime.datetime.now()
        for metric in self.metrics:
            self.data[metric][now] = self.measure(metric)

    def measure(self, metric, device=None):
        if device is None:
            device = smi.listDevices(showall=False)[0]
        if smi is None:
            return 0.0
        elif metric == 'fan_speed_percent':
            return smi.getFanSpeed(device)[1]
        elif metric.find('clk') >=0 and metric.split('_')[0] in smi.validClockNames:
            return int(smi.getCurrentClock(device, metric.split('_')[0], 'freq').strip('Mhz'))
        elif 'used_memory_percent':
            used_bytes, total_bytes = smi.getMemInfo(device, 'vram')
            return int(used_bytes)*100.0/int(total_bytes)
        else:
            raise ValueError('Unrecognized metric requested: {}'.format(metric))

    def save(self, info_filename):
        with open(info_filename, 'w') as output_file:
            output_file.write('# Time, {}\n'.format(', '.join(self.metrics)))
            for time_measurement in sorted(self.data[self.metrics[0]].keys()):
                output_file.write('{}, {}\n'.format(str(time_measurement),
                        ', '.join(str(self.data[metric][time_measurement]) for metric in self.metrics)))
            output_file.close()

    @classmethod
    def from_file(cls, info_filename):
        if pd is None:
            print('WARNING - pandas is required for background system monitor')
            return None
        rv = cls()
        rv.data = pd.read_csv(info_filename, index_col=0, squeeze=True, parse_dates=True).to_dict()
        rv.metrics = [key for key in rv.data.keys()]
        return rv

    def extend(self, other):
        if self.metrics != other.metrics:
            raise ValueError('Both SystemMonitors must have the same record metrics')
        for metric in self.metrics:
            for time_measurement, value in other.data[metric].items():
                self.data[metric][time_measurement] = value

    def get_times(self):
        return self.data[self.metrics[0]].keys()

    def get_start_time(self):
        return min(self.get_times())

    def get_end_time(self):
        return max(self.get_times())

    def plot(self):
        if plt is not None:
            figure, axes = plt.subplots(len(self.metrics), 1, sharex=True, squeeze=False)
            for ax_idx, metric in enumerate(self.metrics):
                ax = axes[ax_idx, 0]
                x_values = sorted(self.data[metric].keys())
                y_values = [self.data[metric][x] for x in x_values]
                ax.plot(x_values, y_values, '.')
                ax.set_ylabel(metric, rotation=0)
            plt.show()

class ArgumentABC(object):
    def __init__(self):
        self._value = None

    def get_args(self):
        raise NotImplementedError('ArgumentABC.get_args is meant to be pure virtual')

    def get_hash(self):
        return '_'.join(self.get_args())

    def get_value(self):
        if self._value is None:
            raise RuntimeError('No value specified! ArgumentABC.get_value can only be used if a value is explicitely set.')
        return self._value

    def set(self, value):
        self._value = value

    def is_set(self):
        return self._value is not None

    def is_shell_only(self):
        '''Returns True if the command must be run with subprocess: shell=True'''
        return False

class PositionalArgument(ArgumentABC):
    def get_args(self):
        if self._value is None:
            raise RuntimeError('No value set for positional argument')
        return [str(self._value)]

class RequiredArgument(ArgumentABC):
    def __init__(self, flag):
        ArgumentABC.__init__(self)
        self.flag = flag

    def get_args(self):
        if self._value is None:
            raise RuntimeError('No value set for {}'.format(self.flag))
        return [self.flag, str(self._value)]

class DefaultArgument(ArgumentABC):
    def __init__(self, flag, default):
        ArgumentABC.__init__(self)
        self.flag = flag
        self.default = default

    def get_args(self):
        if self._value is None:
            return [self.flag, str(self.default)]
        return [self.flag, str(self._value)]

class RepeatedArgument(ArgumentABC):
    def __init__(self, flag):
        ArgumentABC.__init__(self)
        self.flag = flag

    def get_args(self):
        rv = []
        if self._value is None:
            raise RuntimeError('No value set for {}'.format(self.flag))
        for item in self._value:
            rv.extend([self.flag, str(item)])
        return rv

class OptionalArgument(RequiredArgument):
    def get_args(self):
        if self._value is None:
            return []
        return [self.flag, str(self._value)]

class OptionalFlagArgument(ArgumentABC):
    def __init__(self, flag, default = False, add_flag_on_true = True):
        '''Adds `flag` if the `value` is set to `add_flag_on_true` '''
        ArgumentABC.__init__(self)
        self.flag = flag
        self._value = default
        self.add_flag_on_true = add_flag_on_true

    def get_args(self):
        if self._value == self.add_flag_on_true:
            return [self.flag]
        return []

class PipeToArgument(ArgumentABC):
    def get_args(self):
        if self._value is None:
            raise RuntimeError('No value set for pipe to argument')
        return ['2>&1', '|', 'tee', str(self._value)]

    def is_shell_only(self):
        return True

class ExecutionInfo(object):
    def __init__(self, filename):
        self.filename = filename
        self._props = {}
        if os.path.exists(self.filename):
            self._props = json.load(open(self.filename, 'r'))

    def save(self):
        json.dump(self._props, open(self.filename, 'w'), sort_keys=True, indent=4)

    def set_return_code(self, return_code):
        self._props['return_code'] = return_code

    def get_return_code(self):
        return self._props['return_code'] if 'return_code' in self._props else None


class ArgumentSetABC(object):
    def _define_consistent_arguments(self):
        '''Fill self.consistent_args with instances of ArgumentABC.'''
        raise NotImplementedError('ArgumentSetABC._define_consistent_arguments is meant to be pure virtual')

    def _define_variable_arguments(self):
        '''Fill self.variable_args with instances of ArgumentABC.'''
        raise NotImplementedError('ArgumentSetABC._define_variable_arguments is meant to be pure virtual')

    def get_full_command(self, run_configuration):
        '''Translate an instance of RunConfiguration into the full set of command line arguments.'''
        raise NotImplementedError('ArgumentSetABC.get_full_command is meant to be pure virtual')

    def get_interleaved_command(self, run_configurations):
        '''Translate all of the instances of RunConfiguration into a single set of command line arguments
           that generates all output folders in single call.'''
        raise NotImplementedError('ArgumentSetABC.get_interleaved_command is meant to be pure virtual')

    def collect_timing(self, run_configuration):
        '''Use a RunConfiguration to find the data files on disk and process them.'''
        raise NotImplementedError('ArgumentSetABC.collect_timing is meant to be pure virtual')

    def __init__(self, combine_executables=False, **kwargs):
        self.combine_executables = combine_executables
        self.consistent_args = OrderedDict()
        self._define_consistent_arguments()
        self.variable_args = OrderedDict()
        self._define_variable_arguments()
        for key in kwargs:
            self.set(key, kwargs[key])

    def set_user_args(self, user_args):
        ''' Set the command line arguments specified by the user through argparse.
        Not to be confused with the command line arguments that are used to run a benchmark tool.
        Argparse arguments are available to control progam flow, but only after the constructor
        because otherwise child classes would need to correctly pass them in. This function is
        called when argument sets are added to a CommandRunner instance'''
        self.user_args = user_args

    def get_deep_copy(self):
        return copy.deepcopy(self)

    def is_shell_only(self):
        '''Returns True if the command must be run with subprocess: shell=True'''
        for key in self.consistent_args:
            if self.consistent_args[key].is_shell_only():
                return True
        for key in self.variable_args:
            if self.variable_args[key].is_shell_only():
                return True
        return False

    def set(self, key, value):
        if key in self.consistent_args:
            self.consistent_args[key].set(value)
        elif key in self.variable_args:
            self.variable_args[key].set(value)
        else:
            raise ValueError('{} is not a defined argument'.format(key))
        # Add a convience accessor, prefixed with an underscore to denote that it is private/read only.
        self.__setattr__('_' + key, value)

    def set_many(self, kvpairs):
        for key in kvpairs:
            self.set(key, kvpairs[key])

    def get(self, key):
        if key in self.consistent_args:
            return self.consistent_args[key]
        elif key in self.variable_args:
            return self.variable_args[key]
        else:
            raise ValueError('{} is not a defined argument'.format(key))

    def get_args(self, consistent_only=False, ignore_keys=[], require_keys=None):
        rv = []
        for key in self.consistent_args:
            if not key in ignore_keys:
                if require_keys is None or key in require_keys:
                    rv += self.consistent_args[key].get_args()
        if not consistent_only:
            for key in self.variable_args:
                if not key in ignore_keys:
                    if require_keys is None or key in require_keys:
                        rv += self.variable_args[key].get_args()
        return rv

    def __repr__(self):
        arg_values = ['{}:{}'.format(key, self.consistent_args[key]._value) for key in self.consistent_args if self.consistent_args[key]._value is not None]
        return 'ArgumentSet(' + ' '.join(arg_values) + ')'

    # Use this hash of the arguments to remove equivalent runs from the global set of runs
    # Additional constraints on the keys used for the hash can be added for sorting purposes
    def get_hash(self, *args, **kwargs):
        return str(hashlib.md5(' '.join(self.get_args(True, *args, **kwargs)).encode()).hexdigest())

    def get_name(self):
        return '"{}"'.format(' '.join(self.get_args(True)))

    def get_output_basename(self):
        '''Returns a hash of the argument set to create a unique name for the output data.
        Can be overridden to a more intuitive name as long as the returned string is unique
        for a given set of arguments.'''
        return self.get_hash() + '.dat'

    def get_output_subdirectory(self, run_configuration, create=True):
        '''Returns a hash of the argument set to create a unique name for the output data.
        Can be overridden to a more intuitive name as long as the returned string is unique
        for a given set of arguments.'''
        rv = os.path.join(run_configuration.output_directory, self.get_hash())
        if create and not os.path.exists(rv):
            os.makedirs(rv)
        return rv

    def get_output_file(self, run_configuration):
        return os.path.join(run_configuration.output_directory, self.get_output_basename())

    def get_caption(self, similar_keys):
        '''Override this function to make a more meaninful caption based off a subset of keys.'''
        return None

    def _get_stdout_filename(self, run_configuration):
            basename = os.path.splitext(self.get_output_basename())[0]
            return os.path.abspath(os.path.join(run_configuration.output_directory, basename + '.out'))
    def _get_stderr_filename(self, run_configuration):
            basename = os.path.splitext(self.get_output_basename())[0]
            return os.path.abspath(os.path.join(run_configuration.output_directory, basename + '.err'))
    def _get_exec_info_filename(self, run_configuration):
            basename = os.path.splitext(self.get_output_basename())[0]
            return os.path.abspath(os.path.join(run_configuration.output_directory, basename + '.json'))
    def _get_system_monitor_filename(self, run_configuration):
            basename = os.path.splitext(self.get_output_basename())[0]
            return os.path.abspath(os.path.join(run_configuration.output_directory, basename + '.info'))

    def get_system_monitor(self, run_configuration):
        import_rocm_smi(self.user_args.install_path)
        info_filename = self._get_system_monitor_filename(run_configuration)
        return SystemMonitor.from_file(info_filename) if os.path.exists(info_filename) else None

    def execute(self,
                run_configuration = None,
                run_configurations = None,
                overwrite = True,
                dry_run = False):
        if self.combine_executables and (run_configurations is None or run_configuration is not None):
            raise ValueError('A list of run configurations must be passed in when using combined executables!')
        if not self.combine_executables and (run_configuration is None or run_configurations is not None):
            raise ValueError('A single run configuration must be passed in when using individual executables!')
        basename = os.path.splitext(self.get_output_basename())[0]

        # If running multiple, base commands off the first run configuration, and copy output files for others.
        # If running single, create a length 1 list of run configurations to help share code.
        if self.combine_executables:
            run_configuration = run_configurations[0]
        else:
            run_configurations = [run_configuration]

        execution_info = ExecutionInfo(filename = self._get_exec_info_filename(run_configuration))
        old_return_code = execution_info.get_return_code() if not dry_run else None
        if old_return_code is not None and not overwrite:
            message = '{0} Using existing result with code {1} {0}'.format('=' * 10, old_return_code)
            return_code = old_return_code
        else:
            cmd = self.get_interleaved_command(run_configurations) if self.combine_executables else self.get_full_command(run_configuration)
            cmd_str = ' '.join(cmd)
            print(cmd_str)

            if dry_run:
                return_code = 0
            else:
                stdout_file = open(self._get_stdout_filename(run_configuration), mode='w')
                stderr_file = open(self._get_stderr_filename(run_configuration), mode='w')
                # Log some information about the time and command being executed
                time_str = str(datetime.datetime.now())
                for out_file in [stdout_file, stderr_file]:
                    out_file.write('{0} {1} {0}\n'.format('=' * 10, time_str))
                    out_file.write(cmd_str + '\n')
                    out_file.flush()

                import_rocm_smi(self.user_args.install_path)
                system_monitor = SystemMonitor()

                is_shell_only = self.is_shell_only()
                if is_shell_only:
                    cmd = cmd_str
                proc = subprocess.Popen(cmd, stdout=stdout_file, stderr=stderr_file,
                                        env=os.environ.copy(), shell=is_shell_only)
                # Monitor system metrics while the process executes
                poll_metric_count = 0
                try:
                    while proc.poll() is None:
                        if smi is not None and poll_metric_count % 20 == 0:
                            system_monitor.record_line()
                        time.sleep(0.01)
                        poll_metric_count += 1
                except Exception as e:
                    proc.kill()
                    raise(e)

                # Process has completed, collect the return code
                return_code = proc.poll() # return code of process
                execution_info.set_return_code(return_code)
                execution_info.save()
                system_monitor.save(self._get_system_monitor_filename(run_configuration))
                message = '{0} Completed with code {1} {0}'.format('=' * 10, return_code)

                for out_file in [stdout_file, stderr_file]:
                    out_file.write(message + '\n')
                    out_file.flush()

                # Copy output files for each run configuration that is not the first.
                for added_run_configuration in run_configurations[1:]:
                    for filename_fn in [self._get_stdout_filename, self._get_stderr_filename, self._get_exec_info_filename]:
                        shutil.copyfile(filename_fn(run_configuration), filename_fn(added_run_configuration))

        if return_code != 0 or old_return_code is not None:
            print(message)
        return return_code

class ArgumentSetDifference(object):
    def __init__(self, argument_sets, ignore_keys = []):
        if len(argument_sets) <= 0:
            raise ValueError('Expected more than one set of arguments')
        for argument_set in argument_sets:
            if not isinstance(argument_set, ArgumentSetABC):
                raise ValueError('Inputs must derive from ArgumentSetABC')
        self.argument_sets = argument_sets
        self.ignore_keys = ignore_keys
        # Cache a list of differences and similarities with respect to the first input
        self.similarities = []
        self.differences = []
        self.base_argument_set = self.argument_sets[0]
        self.compare_list = self.argument_sets[1:]
        for key in self.base_argument_set.consistent_args:
            if not key in ignore_keys:
                is_similar = True
                for compare_argument_set in self.compare_list:
                    if self.base_argument_set.get(key).get_args() != compare_argument_set.get(key).get_args():
                        is_similar = False
                append_list = self.similarities if is_similar else self.differences
                append_list.append(key)

    def get_differences(self):
        return self.differences

    def get_similarities(self):
        return self.similarities

    def get_as_caption(self):
        custom_caption = self.base_argument_set.get_caption(self.similarities)
        if custom_caption is not None:
            return custom_caption
        rv = 'Constants: '
        for key in self.similarities:
            rv = ' '.join([rv] + self.base_argument_set.get(key).get_args())
        if len(self.differences) > 0:
            rv += '; Differences: '
            for key in self.differences:
                rv += str([' '.join(argument_set.get(key).get_args()) for argument_set in self.argument_sets])
        return rv

class ArgumentSetSort(OrderedDict):
    '''Subclass of OrderedDict that divides a list of argument_sets according to common keys.'''
    def __init__(self, argument_sets, isolate_keys):
        OrderedDict.__init__(self)
        alphabet = [x for x in string.ascii_lowercase]
        if len(alphabet) < len(argument_sets):
            alphabet.extend([x1 + x2 for x1,x2 in itertools.product(string.ascii_lowercase, string.ascii_lowercase)])
        hash_to_label = {}
        alphabet_idx = 0
        for argument_set in argument_sets:
            hash_ignoring = argument_set.get_hash(ignore_keys=isolate_keys)
            if not hash_ignoring in hash_to_label:
                label = 'Run {})'.format(alphabet[alphabet_idx])
                hash_to_label[hash_ignoring] = label
                alphabet_idx += 1
                self[label] = []
            self[hash_to_label[hash_ignoring]].append(argument_set)

class MachineSpecs(dict):
    @classmethod
    def collect_specs(cls, device_numbers, install_path):
        # Helper to translate bytes into human readable units
        def to_mem_units(num_bytes):
            num_bytes = int(num_bytes)
            for exponent, unit in enumerate(['bytes', 'kB', 'MB', 'GB', 'TB']):
                divisor = 1024.0**exponent
                if num_bytes / divisor < 1024.0:
                    break
            return '{:.1f}{}'.format(num_bytes / divisor, unit)
        rv = cls()
        host_info = {}
        host_info['hostname'] = getspecs.gethostname()
        host_info['cpu info'] = getspecs.getcpu()
        host_info['ram'] = getspecs.getram()
        host_info['distro'] = getspecs.getdistro()
        host_info['kernel version'] = getspecs.getkernel()
        host_info['rocm version'] = getspecs.getrocmversion()
        rv['Host'] = host_info

        for device_num in device_numbers:
            device_info = {}
            device_info['device'] = getspecs.getdeviceinfo(device_num)
            device_info['vbios version'] = getspecs.getvbios(device_num)
            device_info['vram'] = getspecs.getvram(device_num)
            device_info['performance level'] = getspecs.getperflevel(device_num)
            device_info['system clock'] = getspecs.getsclk(device_num)
            device_info['memory clock'] = getspecs.getmclk(device_num)
            rv['Device {0:2d}'.format(device_num)] = device_info

        smi = import_rocm_smi(install_path)
        if smi is not None:
            devices = smi.listDevices(showall=False)
            for device in devices:
                smi_info = {}
                smi_info['Bus'] = smi.getBus(device)
                smi_info['Profile'] = smi.getProfile(device)
                smi_info['Start Fan Speed'] = str(smi.getFanSpeed(device)[1]) + '%'
                for clock in smi.validClockNames:
                    freq = smi.getCurrentClock(device, clock, 'freq')
                    measured_level = smi.getCurrentClock(device, clock, 'level')
                    max_level = smi.getMaxLevel(device, clock)
                    smi_info['Start ' + clock] = '{} - Level {}/{}'.format(freq, measured_level, max_level)
                for mem_type in smi.validMemTypes:
                    key = 'Start {} Memory'.format(mem_type)
                    used_bytes, total_bytes = smi.getMemInfo(device, mem_type)
                    smi_info[key] = '{} / {}'.format(to_mem_units(used_bytes), to_mem_units(total_bytes))
                for component in smi.validVersionComponents:
                    smi_info[component.capitalize() + ' Version'] = smi.getVersion([device], component)
                rv['ROCm ' + device.capitalize()] = smi_info

        return rv

    @classmethod
    def from_file(cls, filename):
        rv = cls(json.load(open(filename, 'r')))
        return rv

    def save(self, filename):
        json.dump(self, open(filename, 'w'), sort_keys=True, indent=4)

    def write_latex(self, latex_module):
        for section_key in sorted(self.keys()):
            with latex_module.create(pylatex.FlushLeft()) as centered:
                with centered.create(pylatex.Tabu('ll')) as data_table:
                    header_row = [section_key + ' Info', '']
                    data_table.add_row(header_row, mapper=[pylatex.utils.bold])
                    section_info = self[section_key]
                    for spec_key in sorted(section_info.keys()):
                        data_table.add_row([spec_key + ':', section_info[spec_key]])
            #latex_module.append('\n\n')

    def write_docx(self, document, table_style, level=1):
        for section_key in sorted(self.keys()):
            num_columns = 2
            section_info = self[section_key]
            num_rows = len(section_info) + 1
            document.add_heading(section_key + ' Specifications', level=level)
            table = document.add_table(num_rows, num_columns, style=table_style)
            table.cell(0,0).text = 'Description'
            table.cell(0,1).text = 'Value'
            for row_idx, spec_key in enumerate(sorted(section_info.keys())):
                table.cell(row_idx+1,0).text = str(spec_key)
                table.cell(row_idx+1,1).text = str(section_info[spec_key])

class RunConfiguration(object):
    '''A RunConfiguration contains all of the information that is unique to a set of comparable commands.

    Works in conjunction with ArgumentSetABC to define the complete set of parameters for running an executable.
    ArgumentSetABC should define all of the constant parameters, whereas RunConfiguration defines all of the
    parameters that are being compared between runs. For example, the two required arguments are the
    location of the executable to be tested, and the output directory for the results. To add additional
    comparables, such as the number of GPUs used, derive from this class and add the desired variables.

    An instance of RunConfiguration is passed into ArgumentSetABC.get_full_command. That is where the
    information stored in this class is translated into actual commandline arguments.
    '''
    def __init__(self, user_args, executable_directory, output_directory, label, run_number = None):
        self.user_args = user_args
        self.executable_directory = executable_directory
        self.output_directory = output_directory
        self.label = label
        if run_number is not None:
            self.output_directory = os.path.join(output_directory, 'run{0:02d}'.format(run_number))
        self.run_number = run_number

    def get_hash(self):
        # Assume the internal ouput directory is always unique (because it includes run number)
        return str(hashlib.md5(self.output_directory.encode()).hexdigest())

    def get_id(self):
        # Assume groupable sets of runs have the same label. This has the side-effect of
        # combining different output folders if they use the same label.
        # Consider that effect a feature, not a bug :).
        return self.label

    def make_output_directory(self):
        if not os.path.exists(self.output_directory):
            os.makedirs(self.output_directory)

    def assert_exists(self):
        if not os.path.exists(self.output_directory):
            raise ValueError('Unable to find output directory: {}'.format(self.output_directory))

    def _machine_specs_filename(self):
        return os.path.join(self.output_directory, "specs.json")

    def save_specifications(self, device_num):
        filename = self._machine_specs_filename()
        MachineSpecs.collect_specs([device_num], self.user_args.install_path).save(filename)
        # Does not return the specs because to use them robustly, they need to be loaded
        # from disk. Collecting could overwrite saved specs when post-processing results.

    def load_specifications(self):
        return MachineSpecs.from_file(self._machine_specs_filename())

class RunConfigurationsList(list):
    def group_by_label(self):
        sorted_configurations = OrderedDict()
        for run_configuration in self:
            key = run_configuration.get_id()
            if not key in sorted_configurations:
                sorted_configurations[key] = []
            sorted_configurations[key].append(run_configuration)
        return sorted_configurations

class Comparison(object):
    def __init__(self, argument_sets = [], filename = None, description = None):
        self.argument_sets = copy.deepcopy(argument_sets)
        self.filename = filename
        self.description = description
        self._check_consistency()

    def add(self, argument_set):
        self.argument_sets.append(argument_set)
        self._check_consistency()

    def _check_consistency(self):
        argument_set_hashes = [argument_set.get_hash() for argument_set in self.argument_sets]
        if len(argument_set_hashes) != len(set(argument_set_hashes)):
            raise RuntimeError('Not all run argument sets have a unique hash!')

    def set_user_args(self, user_args):
        ''' Set the command line arguments specified by the user through argparse.
        Not to be confused with the command line arguments that are used to run a benchmark tool.
        Argparse arguments are available to control progam flow, but only after the constructor
        because otherwise child classes would need to correctly pass them in. This function is
        called when comparisons are added to a CommandRunner instance'''
        self.user_args = user_args

    def get_name(self):
        if self.filename is not None:
            return self.filename
        return self.get_hash()

    def get_hash(self):
        # The same set of arguments, but with a different class constitutes a different comparison
        hash = str(self.__class__.__name__)
        for argument_set in self.argument_sets:
            hash += argument_set.get_hash()
        return str(hashlib.md5(hash.encode()).hexdigest())

    def get_caption(self, run_configurations = None):
        rv = self.description if self.description is not None else ''
        if run_configurations is None:
            rv += ' ' + ArgumentSetDifference(self.argument_sets).get_as_caption()
        else:
            grouped_run_configurations = run_configurations.group_by_label()
            if len(grouped_run_configurations) != len(run_configurations):
                rv += ' Bar chart represents median value from multiple runs and the sorted raw data is super-imposed with black markers.'
        if not rv: # Still blank, attempt to collect from ArgumentSerDiffernce
            rv += ArgumentSetDifference(self.argument_sets).get_as_caption()
        return rv

    def _get_sweep_keys(self):
        '''The keys that are collapsed when collecting results. E.g. Used to make the x-axis of a plot.'''
        return []

    def write_latex_table(self, latex_module):
        if len(self.argument_sets) > 0:
            argument_diff = ArgumentSetDifference(self.argument_sets, ignore_keys=self._get_sweep_keys())
            differences = argument_diff.get_differences()
            is_a_comparison = len(differences) > 0
            latex_module.append(
                 ('For all runs, ``' if is_a_comparison else 'Command: ')
                + ' '.join(self.argument_sets[0].get_args(require_keys=argument_diff.get_similarities()))
                +("'' is held constant." if is_a_comparison else '')
                )
            if is_a_comparison:
                with latex_module.create(pylatex.Center()) as centered:
                    tabu_format = 'r|' + ''.join(['c' for key in differences])
                    with centered.create(pylatex.Tabu(tabu_format)) as data_table:
                        header_row = ['label'] + differences
                        data_table.add_row(header_row, mapper=[pylatex.utils.bold])
                        data_table.add_hline()
                        sorted_argument_sets = self.sort_argument_sets(isolate_keys=self._get_sweep_keys())
                        for argument_set_hash, argument_sets in sorted_argument_sets.items():
                            if len(argument_sets) > 0:
                                argument_set = argument_sets[0]
                                row = [argument_set_hash]
                                for key in differences:
                                    argument = argument_set.get(key)
                                    row.append(argument.get_value() if argument.is_set() else 'DEFAULT')
                                data_table.add_row(row)

    def write_docx_table(self, document):
        if len(self.argument_sets) > 0:
            argument_diff = ArgumentSetDifference(self.argument_sets, ignore_keys=self._get_sweep_keys())
            differences = argument_diff.get_differences()
            is_a_comparison = len(differences) > 0
            document.add_paragraph(
                 ('For all runs, ``' if is_a_comparison else 'Command: ')
                + ' '.join(self.argument_sets[0].get_args(require_keys=argument_diff.get_similarities()))
                +("'' is held constant." if is_a_comparison else '')
                )
            if is_a_comparison:
                header_row = ['label'] + differences
                num_columns = len(header_row)
                sorted_argument_sets = self.sort_argument_sets(isolate_keys=self._get_sweep_keys())
                num_rows = len(sorted_argument_sets) + 1
                table_style = 'Colorful Grid' if self.user_args.docx_template is None else None
                table = document.add_table(num_rows, num_columns, style=table_style)
                row_idx = 0
                for col_idx, data in enumerate(header_row):
                    table.cell(row_idx, col_idx).text = data
                for argument_set_hash, argument_sets in sorted_argument_sets.items():
                    if len(argument_sets) > 0:
                        row_idx += 1
                        argument_set = argument_sets[0]
                        row = [argument_set_hash]
                        for key in differences:
                            argument = argument_set.get(key)
                            row.append(argument.get_value() if argument.is_set() else 'DEFAULT')
                        for col_idx, data in enumerate(row):
                            table.cell(row_idx, col_idx).text = str(data)

    def sort_argument_sets(self, isolate_keys):
        return ArgumentSetSort(self.argument_sets, isolate_keys)

    def plot(self, run_configurations, axes):
        num_argument_sets = len(self.argument_sets)
        if num_argument_sets == 0:
            return

        sorted_argument_sets = self.sort_argument_sets(isolate_keys=[]) # No sort applied, but labels provided
        grouped_run_configurations = run_configurations.group_by_label()

        num_groups = len(grouped_run_configurations)
        metric_labels = [key for key in self.argument_sets[0].collect_timing(run_configurations[0])]
        num_metrics = len(metric_labels)
        if num_metrics == 0:
            return
        num_comparables = num_groups * num_argument_sets
        x_locations = np.arange(num_metrics)
        width = 0.8 / num_comparables
        offset_start = -0.4 + width/2.0
        label_map = OrderedDict()

        # Create a mapping of x-locations for the bar chart.
        x_mapping = {}
        gap_scalar = 0.9
        for x_metric_offset, metric_label in zip(x_locations, metric_labels):
            x_mapping[metric_label] = {}
            cmp_idx = 0
            for subset_label in sorted_argument_sets:
                x_mapping[metric_label][subset_label] = {}
                group_linear_mapping = []
                for group_label in grouped_run_configurations:
                    group_linear_mapping.append(
                            x_metric_offset + offset_start + width*cmp_idx
                            )
                    cmp_idx += 1
                # scale down the size of each subset plot to 90% to create a gap
                group_linear_mapping = np.array(group_linear_mapping)
                group_linear_mapping_mean = np.mean(group_linear_mapping)
                group_linear_mapping -= group_linear_mapping_mean
                group_linear_mapping *= gap_scalar
                group_linear_mapping += group_linear_mapping_mean
                for group_idx, group_label in enumerate(grouped_run_configurations):
                    x_mapping[metric_label][subset_label][group_label] = (
                            group_linear_mapping[group_idx]
                            )

        def map_to_x_list(subset_label, metric_label, group_label, y_list = [0, ]):
            num_x = len(y_list)
            sub_x_offsets = np.linspace(-0.4*width, 0.4*width, num_x+1, endpoint=True)
            return [
                x_mapping[metric_label][subset_label][group_label]
                + 0.5 *(sub_x_offsets[data_idx] + sub_x_offsets[data_idx+1])
                for data_idx in range(num_x)]

        # loop over independent outputs
        x_scatter_by_group = OrderedDict()
        y_scatter_by_group = OrderedDict()
        x_bar_by_group = OrderedDict()
        y_bar_by_group = OrderedDict()
        for group_label, run_configuration_group in grouped_run_configurations.items():
            x_scatter_by_group[group_label] = []
            y_scatter_by_group[group_label] = []
            x_bar_by_group[group_label] = []
            y_bar_by_group[group_label] = []
            # loop over argument sets that differ other than the swept variable(s)
            for subset_label, partial_argument_sets in sorted_argument_sets.items():
                if len(partial_argument_sets) != 1:
                    raise ValueError('Assumed that sorting argument sets with no keys has a single element per sort.')
                argument_set = partial_argument_sets[0]
                y_list_by_metric = OrderedDict() # One array of y values for each metric
                # loop over number of coarse grain runs and concatenate results
                for run_configuration in run_configuration_group:
                    timing_results = argument_set.collect_timing(run_configuration)
                    for metric_label in timing_results:
                        if not metric_label in y_list_by_metric:
                            y_list_by_metric[metric_label] = []
                        y_list_by_metric[metric_label].extend(timing_results[metric_label])
                # For each metric, add a set of bars in the bar chart.
                for metric_label, y_list in y_list_by_metric.items():
                    x_list = map_to_x_list(subset_label, metric_label, group_label, y_list)
                    x_scatter_by_group[group_label].extend(x_list)
                    y_scatter_by_group[group_label].extend(sorted(y_list))
                    x_bar_by_group[group_label].append(np.mean(x_list))
                    y_bar_by_group[group_label].append(np.median(y_list))
        for group_label in x_scatter_by_group:
            axes.bar(
                    x_bar_by_group[group_label],
                    y_bar_by_group[group_label],
                    gap_scalar * width,
                    label = group_label,
                    )
            axes.plot(
                    x_scatter_by_group[group_label],
                    y_scatter_by_group[group_label],
                    'k*',
                    )
        # If a single metric is provided, it is labelled on the y-axis.
        # If multiple metrics are provided, they are labelled along the x-axis and
        # then the units are assumed to be Time (s).
        if len(metric_labels) != 1:
            axes.set_xticks(x_locations)
            axes.set_xticklabels(metric_labels)
        else:
            axes.set_xticks([])
        # Set the minor tick labels as a legend
        minor_ticks = []
        minor_labels = []
        for metric_label in metric_labels:
            for subset_label in sorted_argument_sets:
                group_x = []
                for group_label in grouped_run_configurations:
                    group_x.extend(map_to_x_list(subset_label, metric_label, group_label))
                minor_ticks.append(np.mean(group_x))
                minor_labels.append(subset_label)
        axes.xaxis.set_minor_formatter(FuncFormatter(lambda x, pos : str(x)))
        axes.set_xticks(minor_ticks, minor=True)
        axes.set_xticklabels(minor_labels, minor=True)

        axes.set_xlim(x_locations[0] - 0.5, x_locations[-1] + 0.5)
        axes.set_ylabel(metric_labels[0] if len(metric_labels) == 1 else 'Time (s)' )
        return True

    def custom_plot(self, run_configurations, is_make_plot):
        plot_filename = None
        plot_caption = None
        return plot_filename, plot_caption

class SingleCommand(object):
    def __init__(self, argument_set, run_configuration):
        self.argument_set = argument_set
        self.run_configuration = run_configuration

    def execute(self, **kwargs):
        self.argument_set.execute(run_configuration = self.run_configuration, **kwargs)

class CombinedCommand(object):
    def __init__(self, argument_set, run_configurations):
        self.argument_set = argument_set
        self.run_configurations = run_configurations

    def execute(self, **kwargs):
        self.argument_set.execute(run_configurations = self.run_configurations, **kwargs)

class CommandList(object):
    def __init__(self):
        self.commands = []

    def add_command(self, argument_set, run_configurations):
        if argument_set.combine_executables:
            self.commands.append(CombinedCommand(argument_set, run_configurations))
        else:
            for run_configuration in run_configurations:
                self.commands.append(SingleCommand(argument_set, run_configuration))

    def execute_shuffled(self, **kwargs):
        random.Random(8341).shuffle(self.commands) # Randomize, but with a consistent seed.
        num_commands = len(self.commands)
        for idx, command in enumerate(self.commands):
            print('Running command {} of {}'.format(idx+1, num_commands))
            command.execute(**kwargs)

class CommandRunner(object):
    def __init__(self, user_args, run_configuration_cls = RunConfiguration):
        self.user_args = user_args

        executable_directories = user_args.input_executables
        output_directories = user_args.output_directories
        labels = user_args.labels

        print('Excecutable directories: ', executable_directories)

        if len(executable_directories) > len(output_directories):
            for i in range(len(output_directories), len(executable_directories)):
                output_directories.append('dir' + str(i))
        print('Output directories: ', output_directories)

        if len(output_directories) > len(labels):
            for i in range(len(labels), len(output_directories)):
                labels.append(os.path.basename(output_directories[i].strip('/')))
        print('Run labels:', labels)

        print('Document output: ', user_args.documentation_directory)
        if not os.path.exists(user_args.documentation_directory):
            os.makedirs(user_args.documentation_directory)

        self.executable_directories = executable_directories
        self.output_directories = output_directories
        self.labels = labels

        self.run_configurations = RunConfigurationsList()
        for exec_dir, out_dir, label in zip(executable_directories, output_directories, labels):
            for run_number in range(user_args.num_repititions):
                self.run_configurations.append(run_configuration_cls(
                        user_args = user_args,
                        executable_directory = exec_dir,
                        output_directory = out_dir,
                        label = label,
                        run_number = run_number,
                        ))

        self.argument_set_map = {}
        self.comparison_map = OrderedDict()
        self._check_consistency()

        if self.is_use_pylatex():
            geometry_options = {"margin": "0.7in"}
            self.doc = pylatex.Document(
                os.path.join(self.user_args.documentation_directory, 'latex_summary'),
                geometry_options=geometry_options)

            header = pylatex.PageStyle("header")
            with header.create(pylatex.Foot("L")):
                header.append("AMD Internal Use Only")
            with header.create(pylatex.Foot("R")):
                header.append(pylatex.NoEscape(r'\today'))
            self.doc.preamble.append(header)
            self.doc.change_document_style("header")

            self.doc.preamble.append(pylatex.Command('title', pylatex.NoEscape(r'Benchmark Summary \\ \large AMD Internal Use Only')))
            self.doc.preamble.append(pylatex.Command('author', getpass.getuser()))
            self.doc.preamble.append(pylatex.Command('date', pylatex.NoEscape(r'\today')))
            self.doc.append(pylatex.NoEscape(r'\maketitle'))

        if self.is_use_docx():
            # Author, date and Internal Only, page numbers, etc. set by template file
            self.docx_doc = docx.Document(self.user_args.docx_template)
            if self.user_args.docx_template is None:
                self.docx_doc.add_heading('Benchmark Summary', 0)
                self.docx_doc.add_paragraph('AMD Internal Use Only')
                self.docx_doc.add_paragraph('Default formatting of this auto-generated document is not ideal.'
                                            ' Consider using PDF or supplying a document with a style guide.'
                                            ' Tables and figures will be appended to the end of the input document.')

    def _check_consistency(self):
        run_configuration_hashes = [run_configuration.get_hash() for run_configuration in self.run_configurations]
        if len(run_configuration_hashes) != len(set(run_configuration_hashes)):
            raise RuntimeError('Not all run configurations have a unique hash! Are the output directories unique?')

    def main(self):
        self.execute()
        self.show_plots()
        self.get_system_summary()
        self.output_summary()

    def is_run_tool(self):
        return 'EXECUTE' in self.user_args.methods

    def is_dry_run(self):
        is_dry_run = ('DRY' in self.user_args.methods)
        if self.is_run_tool() and is_dry_run:
            raise ValueError('DRY and EXECUTE are mutually exclusive. Both were specified.')
        return is_dry_run

    def is_make_plots(self):
        return 'PLOT' in self.user_args.methods

    def is_use_matplotlib(self):
        if self.is_make_plots():
            if plt is None:
                print('WARNING - Matplotlib is recommended!')
                return False
            if np is None:
                print('WARNING - Numpy is recommended!')
                return False
            return True
        return False

    def is_make_document(self):
        return 'DOCUMENT' in self.user_args.methods

    def is_use_pylatex(self):
        if self.is_make_document():
            if pylatex is None:
                print('WARNING - PyLaTeX is required for PDF summary!')
                return False
            return True
        return False

    def is_use_docx(self):
        if self.is_make_document():
            if docx is None:
                print('WARNING - docx package is required for .docx summary!')
                return False
            if BytesIO is None:
                print('WARNING - BytesIO package is required for .docx summary!')
                return False
            return True
        return False

    def is_interactive(self):
        return 'INTERACTIVE' in self.user_args.methods

    def is_overwrite(self):
        return 'OVERWRITE' in self.user_args.methods

    def setup_system(self):
        for run_configuration in self.run_configurations:
            if self.is_run_tool():
                run_configuration.make_output_directory()
                run_configuration.save_specifications(self.user_args.device_num)
            elif not self.is_dry_run():
                run_configuration.assert_exists()


    def add_argument_set(self, argument_set):
        argument_set.set_user_args(self.user_args)
        self.argument_set_map[argument_set.get_hash()] = argument_set

    def add_comparisons(self, comparisons):
        for comparison in comparisons:
            comparison_name = comparison.get_name()
            if comparison_name in self.comparison_map:
                print(comparison.argument_sets)
                raise ValueError('Comparison {} was added twice'.format(comparison_name))
            comparison.set_user_args(self.user_args)
            self.comparison_map[comparison_name] = comparison
            for argument_set in comparison.argument_sets:
                self.add_argument_set(argument_set)

    def _filter_argument_set(self, argument_set):
        args = ' '.join(argument_set.get_args(True))
        for required_arg in self.user_args.filter_in:
            if args.find(required_arg) < 0:
                return False
        for banned_arg in self.user_args.filter_out:
            if args.find(banned_arg) >= 0:
                return False
        return True

    def execute(self):
        command_list = CommandList()
        for cmd_hash, argument_set in self.argument_set_map.items():
            if self._filter_argument_set(argument_set):
                command_list.add_command(argument_set, self.run_configurations)
        self.command_list = command_list

        if self.is_run_tool() or self.is_dry_run():
            command_list.execute_shuffled(overwrite = self.is_overwrite(), dry_run = self.is_dry_run())

    def show_plots(self):
        if self.is_dry_run():
            return
        grouped_run_configurations = self.run_configurations.group_by_label()
        for group_label, run_configuration_group in grouped_run_configurations.items():
            run_configuration = run_configuration_group[0]
            machine_specs = run_configuration.load_specifications()
            if self.is_use_pylatex():
                with self.doc.create(pylatex.Section('{} Specifications'.format(group_label))):
                    machine_specs.write_latex(self.doc)
            if self.is_use_docx():
                table_style = 'Light Grid' if self.user_args.docx_template is None else None
                machine_specs.write_docx(self.docx_doc, table_style)
                self.docx_doc.add_page_break()

        active_plots = [] # show plots in batches
        docx_fig_count = 1
        for comparison_name, comparison in self.comparison_map.items():
            # Create any non-matplotlib plots within Comparison.custom_plot()
            plot_filename, plot_caption = comparison.custom_plot(self.run_configurations, self.is_make_plots())
            if plot_filename is not None and os.path.exists(plot_filename):
                if self.is_use_pylatex():
                    with self.doc.create(pylatex.Figure(position='htbp')) as plot:
                        plot.add_image(os.path.abspath(plot_filename), width=pylatex.NoEscape(r'0.8\textwidth'))
                        if plot_caption:
                            plot.add_caption(plot_caption)
                        else:
                            plot.add_caption(comparison.get_caption(self.run_configurations))
                            plot.append(pylatex.NoEscape(r'\vspace{0.3cm}'))
                            comparison.write_latex_table(plot)
                    self.doc.append(pylatex.NoEscape(r'\clearpage'))


            # Add any Matplotlib plots using Comparison.plot()
            if self.is_use_matplotlib():
                figure, axes = plt.subplots(figsize = (7, 7))
                plot_success = comparison.plot(self.run_configurations, axes)
                print(comparison.get_caption(self.run_configurations))
                if plot_success:
                    axes.legend(fontsize = 10, bbox_to_anchor=(0., 1.02, 1., .102), loc='lower left',
                                ncol=2, mode='expand', borderaxespad=0.)
                    figure.tight_layout(rect=(0,0.05,1.0,1.0))

                    if self.is_use_pylatex():
                        with self.doc.create(pylatex.Figure(position='htbp')) as plot:
                            plot.add_plot(width=pylatex.NoEscape(r'0.8\textwidth'), dpi=300, transparent=True)
                            plot.add_caption(comparison.get_caption(self.run_configurations))
                            plot.append(pylatex.NoEscape(r'\vspace{0.3cm}'))
                            comparison.write_latex_table(plot)
                        self.doc.append(pylatex.NoEscape(r'\clearpage'))

                    if self.is_use_docx():
                        memfile = BytesIO()
                        figure.savefig(memfile, format='png', dpi=300, transparent=True)
                        self.docx_doc.add_picture(memfile, width=docx.shared.Inches(6.0))
                        caption_style = 'Quote' if self.user_args.docx_template is None else None
                        self.docx_doc.add_paragraph('Figure {}: '.format(docx_fig_count) + comparison.get_caption(self.run_configurations), style=caption_style)
                        comparison.write_docx_table(self.docx_doc)
                        self.docx_doc.add_page_break()
                        docx_fig_count += 1
                        memfile.close()

                    figure.suptitle(comparison.get_caption(self.run_configurations),
                                    fontsize='medium', y=0.02, va='bottom')
                    figure.savefig(os.path.join(self.user_args.documentation_directory,
                                                comparison.get_name() + '_auto_plot.pdf'))

                if not self.is_interactive():
                    plt.close(figure)
                else:
                    active_plots.append(figure)
                    if len(active_plots) >= 10:
                        plt.show()
                        for open_figure in active_plots:
                            plt.close(open_figure)
                        active_plots = []

        # Show remaining plots if applicable
        if len(active_plots) > 0:
            plt.show()

    def get_system_summary(self):
        if not self.is_interactive():
            return
        total_system_monitor = None
        for cmd_hash, argument_set in self.argument_set_map.items():
            if self._filter_argument_set(argument_set):
                for run_configuration in self.run_configurations:
                    run_system_monitor = argument_set.get_system_monitor(run_configuration)
                    if run_system_monitor is not None:
                        if total_system_monitor is None:
                            total_system_monitor = run_system_monitor
                        else:
                            total_system_monitor.extend(run_system_monitor)
        if total_system_monitor is not None:
            start = total_system_monitor.get_start_time()
            end = total_system_monitor.get_end_time()
            total_time = (end - start).total_seconds()
            print('Test ran from {} to {}. A total of {} seconds.'.format(start, end, total_time))
            total_system_monitor.plot()

    def output_summary(self):
        if self.is_use_pylatex():
            current_working_directory = os.getcwd()
            try:
                self.doc.generate_pdf(clean_tex=False)
            except subprocess.CalledProcessError:
                print('WARNING: Failed to output document')
            #self.doc.generate_tex()
            os.chdir(current_working_directory)
        if self.is_use_docx():
            self.docx_doc.save(os.path.join(self.user_args.documentation_directory, 'benchmark_summary.docx'))

def parse_input_arguments(parser):
    def to_multiple_choices(choices, input_string):
        rv = input_string.upper()
        if not rv in choices:
            raise argparse.ArgumentTypeError('Method must be one of {}. Received {}.'.format(choices, rv))
        return rv

    all_test_methods     = ['DRY', 'EXECUTE', 'OVERWRITE', 'PROCESS', 'PLOT', 'DOCUMENT', 'INTERACTIVE']
    default_test_methods = [       'EXECUTE', 'OVERWRITE', 'PROCESS', 'PLOT', 'DOCUMENT']
    def to_test_methods(s):
        return to_multiple_choices(all_test_methods, s)

    parser.add_argument('-i', '--input-executables', action='append', required=True,
                        help='Input executable location, can be added multiple times.')
    parser.add_argument('-o', '--output-directories', action='append', default=[],
                        help=('Output directories. If more than one input executable is specified,'
                             +' then an output directory must be specified for each.'
                             #+' If a single executable was used for multiple runs, outputs can still be multiply specified.'
                             ))
    parser.add_argument('-l', '--labels', action='append', default=[],
                        help=('Labels for comparing multiple runs. If more than one output is specified,'
                             +' then a label may be specified for each.'
                             +' defaults to the basename of the output directory.'
                             ))
    parser.add_argument('-m', '--methods', default=default_test_methods, nargs='+',
                        type = to_test_methods,
                        help=('Execute a portion of the benchmark suite, or just a subset.'
                             +' Choose a space separated list from [{}]'.format(' '.join(all_test_methods))
                             +' To generate a document without re-running the benchmarks, use `-m PLOT DOCUMENT`.'
                             +' To run without plotting/documentation tools use `-m EXECUTE` and post-process later.'
                             +' To generate plots without creating a summary document use `-m EXECUTE PLOT`.'
                             +' To interact with plots after generating the data use `-m PLOT INTERACTIVE`.'
                             +' By default, existing results are overwritten, but `-m EXECUTE PLOT DOCUMENT` can be used to restart killed runs (omit `OVERWRITE`).'
                             ))
    parser.add_argument('--filter-in', default=[], nargs='+',
                        help=('Space separated list of strings that must be part of the command line string'
                             +' to be considered for this run. Due to limitations in argparse, the input cannot'
                             +' have a leading dash, even when encapsulated with quotes. Note that quotes are'
                             +' likely required because any key-value pair will be separated by a space.'
                             ))
    parser.add_argument('--filter-out', default=[], nargs='+',
                        help=('Space separated list of strings that must NOT be part of the command line string'
                             +' to be considered for this run. Due to limitations in argparse, the input cannot'
                             +' have a leading dash, even when encapsulated with quotes. Note that quotes are'
                             +' likely required because any key-value pair will be separated by a space.'
                             ))
    parser.add_argument('-n', '--num-repititions', default=1, type=int,
                        help='Number of times to run the exectuable.')
    parser.add_argument('--docx-template', default=None,
                        help='Empty docx that contains only a style guide to be used as a template for the summary document.')
    parser.add_argument('-w', '--documentation-directory', default='doc',
                        help='Output directory for the summary documentation.')
    parser.add_argument('-d', '--device-num', default=0, type=int,
                        help='Device number to run on.')
    parser.add_argument('--install-path', default='/opt/rocm', help='Top directory of driver installation.')
    return parser.parse_args()


if __name__ == '__main__':
    import re
    # As an example, profile various modes of the bash function "ls".
    class ListDirArgumentSet(ArgumentSetABC):
        def _define_consistent_arguments(self):
            self.consistent_args['human_readable'] = OptionalFlagArgument('-h')
            self.consistent_args['summarize'] = OptionalFlagArgument('-s')
            self.consistent_args['depth'] = OptionalArgument('-d')

        def _define_variable_arguments(self):
            # Instead of coparing the performance of different executables, we are timing the
            # performance of running the command on the "executable directory"
            self.variable_args['target'] = PositionalArgument()
            self.variable_args['output_file'] = PipeToArgument()

        def __init__(self,
                     human_readable,
                     summarize,
                     depth = None,
                     ):
            # The base constructor could be used directly, but only if all of the inputs
            # are specified as kwargs.
            ArgumentSetABC.__init__(
                    self,
                    human_readable = human_readable,
                    summarize = summarize,
                    depth = depth
                    )

        def get_full_command(self, run_configuration):
            self.set('output_file', self.get_output_file(run_configuration))
            self.set('target', run_configuration.target)
            return ['time', 'du'] + self.get_args()

        def collect_timing(self, run_configuration):
            output_filename = self.get_output_file(run_configuration)
            rv = {}
            if os.path.exists(output_filename):
                timing_file = open(output_filename, 'r')
                output_text = timing_file.read()
                match = re.search(r'(\d+\.\d+)user', output_text)
                if match:
                    rv['timing'] = [float(match.group(1))]
                match = re.search(r'(\d+\.\d+)system', output_text)
                if match:
                    rv['user'] = [float(match.group(1))]
            return rv


    class ListDirRunConfiguration(RunConfiguration):
        def __init__(self, user_args, executable_directory, *args, **kwargs):
            RunConfiguration.__init__(self, user_args, executable_directory, *args, **kwargs)
            self.target = executable_directory

    def create_comparisons():
        # Comparisons are built upon a list of arguments. However, it is acceptable to
        # recreate the same arguments multiple times because each unique set will only
        # be executed once.
        comparisons = []
        for human_readable in [True, False]:
            comparison = Comparison(
                    description = 'Check the cost of the summarize option.',
                    )
            for summarize in [True, False]:
                comparison.add(ListDirArgumentSet(
                        human_readable = human_readable,
                        summarize = summarize,
                        depth = 0,
                        ))
            comparisons.append(comparison)
        for depth in [0, 1, 2]:
            comparison = Comparison(
                    description = 'Fix the depth, and compare the cost of using human readable.',
                    )
            for human_readable in [True, False]:
                comparison.add(ListDirArgumentSet(
                        human_readable = human_readable,
                        summarize = False,
                        depth = depth,
                        ))
            comparisons.append(comparison)
        single_run_comparison = Comparison(
            description = 'Single run without comparisons',
            )
        single_run_comparison.add(ListDirArgumentSet(
                human_readable = True,
                summarize = summarize,
                depth = 2,
                ))
        comparisons.append(single_run_comparison)
        return comparisons

    # location of if __name__ == '__main__': in a normal script
    print('Suggested Usage: python3 commandrunner.py -i / -i /sys -o /tmp/bench_df_root -o /tmp/bench_df_sys'
        + ' -n 3 -m EXECUTE PLOT DOCUMENT')

    # Create a parser and optionally add custom user inputs
    parser = argparse.ArgumentParser()

    # Add the commandrunner specific user inputs and parse them
    user_args = parse_input_arguments(parser)

    command_runner = CommandRunner(user_args, ListDirRunConfiguration)

    command_runner.setup_system()

    command_runner.add_comparisons(create_comparisons())

    command_runner.main()
