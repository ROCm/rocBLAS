#!/usr/bin/python3

import sys, getopt
sys.path.append('../../../clients/common/')
import numpy as np
from math import *
import subprocess
import os
import re # regexp package
import shutil
import tempfile
import rocblas_gentest as gt
import getspecs

usage = '''A timing script for rocBLAS the generates plots

Usage:

\talltime.py
\t\t-A          working directory A
\t\t-B          working directory B (optional)
\t\t-i          input yaml
\t\t-o          output directory
\t\t-b          output directory for the base run
\t\t-T          do not perform BLAS functions; just generate document
\t\t-f          document format: pdf (default) or docx (Need forked docx plugin)
\t\t-d          device number (default: 0)
\t\t-g          generate graphs via Asymptote: 0(default) or 1
\t\t-S          plot speedup (default: 1, disabled: 0)
\t\t-X          do not generate figures
'''

# \t\t-t          data type: gflops #Maybe use option to plot time graphs too


def nextpow(val, radix):
    x = 1
    while(x <= val):
        x *= radix
    return x

class rundata:

    def __init__(self, wdir, odir, diridx, label,
                 data, hwinfo):
        self.wdir = wdir
        self.odir = odir
        self.diridx = diridx
        self.minnsize = data['n']
        self.maxNsize = data['N']
        self.minmsize = data['m']
        self.maxMsize = data['M']
        self.minksize = data['k']
        self.maxKsize = data['K']
        self.nbatch = data['batch_count']
        self.function = data['function']
        self.precision = data['compute_type']    #This picks precision
        self.label = label
        self.incx = data['incx']
        self.incy = data['incy']
        self.alpha = data['alpha']
        self.beta = data['beta']
        self.iters = data['iters']
        self.cold_iters = data['cold_iters']
        self.samples = data['samples']
        self.lda = data['lda']
        self.ldb = data['ldb']
        self.ldc = data['ldc']
        self.LDA = data['LDA']
        self.LDB = data['LDB']
        self.LDC = data['LDC']
        self.algo = data['algo']
        self.transA = data['transA']
        self.transB = data['transB']
        self.initialization = data['initialization']
        self.step_size = data['step_size']
        self.step_mult = data['step_mult']
        self.side = data['side']
        self.uplo = data['uplo']
        self.diag = data['diag']
        self.theoMax = -1


        flopseq = data['flops']
        memeq = data['mem']
        precisionBits = int(re.search(r'\d+', self.precision).group())
        if(self.function == 'trsm' or self.function == 'gemm'):  #TODO better logic to decide memory bound vs compute bound
            self.theoMax = hwinfo['theoMaxCompute'] * 32.00 / precisionBits #scaling to appropriate precision
        elif flopseq and memeq:                                  # Memory bound
            try:
                n=100000
                flops = eval(flopseq)
                mem = eval(memeq)
                self.theoMax = hwinfo['Bandwidth'] / float(eval(memeq)) * eval (flopseq) * 32 / precisionBits / 4   #4 bytes scaled to type
            except:
                print("flops and mem equations produce errors")


    def outfilename(self):
        outfile = str(self.function)
        outfile += "_" + self.precision
        outfile += "_" + self.label.replace(' ', '_').replace('/', '_')
        outfile += ".dat"
        outfile = os.path.join(self.odir, outfile)
        return outfile

    def runcmd(self, nsample):
        cmd = ["./timing.py"]

        cmd.append("-w")
        cmd.append(self.wdir)

        cmd.append("-i")
        cmd.append(str(self.iters))

        cmd.append("-j")
        cmd.append(str(self.cold_iters))

        cmd.append("-a")
        cmd.append(str(self.samples))

        cmd.append("-b")
        cmd.append(str(self.nbatch))

        cmd.append("-m")
        cmd.append(str(self.minmsize))
        cmd.append("-M")
        cmd.append(str(self.maxMsize))

        cmd.append("-n")
        cmd.append(str(self.minnsize))
        cmd.append("-N")
        cmd.append(str(self.maxNsize))

        cmd.append("-k")
        cmd.append(str(self.minksize))
        cmd.append("-K")
        cmd.append(str(self.maxKsize))

        cmd.append("-f")
        cmd.append(self.function)

        cmd.append("-p")
        cmd.append(self.precision)

        cmd.append("--lda")
        cmd.append(str(self.lda))

        cmd.append("--LDA")
        cmd.append(str(self.LDA))

        cmd.append("--ldb")
        cmd.append(str(self.ldb))

        cmd.append("--LDB")
        cmd.append(str(self.LDB))

        cmd.append("--ldc")
        cmd.append(str(self.ldc))

        cmd.append("--LDC")
        cmd.append(str(self.LDC))

        cmd.append("--algo")
        cmd.append(str(self.algo))

        cmd.append("--incx")
        cmd.append(str(self.incx))

        cmd.append("--incy")
        cmd.append(str(self.incy))

        cmd.append("--alpha")
        cmd.append(str(self.alpha))

        cmd.append("--beta")
        cmd.append(str(self.beta))

        cmd.append("--transA")
        cmd.append(str(self.transA))

        cmd.append("--transB")
        cmd.append(str(self.transB))

        cmd.append("--side")
        cmd.append(str(self.side))

        cmd.append("--uplo")
        cmd.append(str(self.uplo))

        cmd.append("--diag")
        cmd.append(str(self.diag))

        cmd.append("-s")
        cmd.append(str(self.step_size))

        if self.step_mult == 1:
            cmd.append("-x")

        cmd.append("-o")
        cmd.append(self.outfilename())

        # cmd.append("-t")
        # cmd.append("gflops")

        return cmd

    def executerun(self, nsample):
        fout = tempfile.TemporaryFile(mode="w+")
        ferr = tempfile.TemporaryFile(mode="w+")

        proc = subprocess.Popen(self.runcmd(nsample), stdout=fout, stderr=ferr,
                                env=os.environ.copy())
        proc.wait()
        rc = proc.returncode
        if rc != 0:
            print("****fail****")
        return rc

class yamldata:

    def __init__(self, configFile):
        self.configFile = configFile
        self.testcases = []
        self.executerun()

    def reorderdata(self):
        oldData = self.testcases
        newData = []
        names = []
        for test in oldData:
            name = test['function']
            precision = test['compute_type']
            side = test['side']
            if (name,precision) not in names:
                type = [ x for x in oldData if x['function']==name and x['compute_type'] == precision and x['side'] == side ]
                newData.append(type)
                names.append((name,precision, side))
        self.testcases = newData

    #Monkey Patch
    def write_test(self, test):
        self.testcases.append(test)

    #Monkey Patch
    def process_doc(self, doc):
        """Process one document in the YAML file"""

        # Ignore empty documents
        if not doc or not doc.get('Tests'):
            return

        # Clear datatypes and params from previous documents
        gt.datatypes.clear()
        gt.param.clear()

        # Return dictionary of all known datatypes
        gt.datatypes.update(gt.get_datatypes(doc))

        # Arguments structure corresponding to C/C++ structure
        gt.param['Arguments'] = type('Arguments', (gt.ctypes.Structure,),
                                {'_fields_': gt.get_arguments(doc)})

        # Special names which get expanded as lists of arguments
        gt.param['dict_lists_to_expand'] = doc.get('Dictionary lists to expand') or ()

        # Lists which are not expanded
        gt.param['lists_to_not_expand'] = doc.get('Lists to not expand') or ()

        # Defaults
        defaults = doc.get('Defaults') or {}

        default_add_ons = {"m": 1, "M": 1, "n": 1, "N": 1, "k": 1, "K": 1, "lda": 1, "ldb": 1, "ldc": 1, "LDA": 1, "LDB": 1, "LDC": 1, "iters": 1, "flops": '', "mem": '', "samples": 1, "step_mult": 0}
        defaults.update(default_add_ons)

        # Known Bugs
        gt.param['known_bugs'] = doc.get('Known bugs') or []

        # Functions
        gt.param['Functions'] = doc.get('Functions') or {}

        # Instantiate all of the tests, starting with defaults
        for test in doc['Tests']:
            case = defaults.copy()
            case.update(test)
            gt.generate(case, gt.instantiate)

    def importdata(self):
        gt.args['includes'] = []
        gt.args['infile'] = self.configFile
        gt.write_test = self.write_test
        for doc in gt.get_yaml_docs():
            self.process_doc(doc)
        # timeCases.extend(self.testcases)
        # print(timeCases)


    def executerun(self):
        self.importdata()
        self.reorderdata()


class getVersion:
    def __init__(self, wdir):
        self.wdir = wdir
        self.prog = os.path.join(self.wdir, "rocblas-bench")
        self.version = ""
        self.executerun()

    def runcmd(self):
        cmd = [self.prog]

        cmd.append("--version")

        return cmd

    def executerun(self):
        fout = os.popen(" ".join(self.runcmd())).read()
        #self.version = fout.split(":",1)[0] + ": " + fout.split("rel-",1)[1]
        self.version = fout

class figure:
    def __init__(self, name, caption):
        self.name = name
        self.runs = []
        self.caption = caption

    def inputfiles(self):
        files = []
        for run in self.runs:
            files.append(run.outfilename())
        return files

    def labels(self):
        labels = []
        for run in self.runs:
            labels.append(run.label)
        return labels

    def filename(self, outdir, docformat):
        outfigure = self.name
        outfigure += ".pdf"
        # if docformat == "pdf":
        #     outfigure += ".pdf"
        # if docformat == "docx":
        #     outfigure += ".png"
        return os.path.join(outdir, outfigure)


    def asycmd(self, outdir, docformat, datatype, ncompare):
        asycmd = ["asy"]

        asycmd.append("-f")
        asycmd.append("pdf")
        # if docformat == "pdf":
        #     asycmd.append("-f")
        #     asycmd.append("pdf")
        # if docformat == "docx":
        #     asycmd.append("-f")
        #     asycmd.append("png")
        #     asycmd.append("-render")
        #     asycmd.append("8")
        asycmd.append("datagraphs.asy")

        asycmd.append("-u")
        asycmd.append('filenames="' + ",".join(self.inputfiles()) + '"')

        asycmd.append("-u")
        asycmd.append('legendlist="' + ",".join(self.labels()) + '"')

        # if dirB != None and speedup: # disabled for now
        #     asycmd.append("-u")
        #     asycmd.append('speedup=true')
        # else:
        #     asycmd.append("-u")
        #     asycmd.append('speedup=false')

        asycmd.append("-u")
        asycmd.append('speedup=' + str(ncompare))

        if datatype == "gflops":
            asycmd.append("-u")
            asycmd.append('ylabel="GFLOP/sec"')

        if self.runs[0].theoMax != -1:
            asycmd.append("-u")
            asycmd.append('theoMax=' + str(self.runs[0].theoMax))

        asycmd.append("-u")
        asycmd.append('xlabel='+'"Size: ' + getXLabel(self.runs[0])+'"')

        asycmd.append("-o")
        asycmd.append(self.filename(outdir, docformat) )

        print(" ".join(asycmd))

        return asycmd

    def executeasy(self, outdir, docformat, datatype, ncompare):
        asyproc =  subprocess.Popen(self.asycmd(outdir, docformat, datatype, ncompare),
                                    env=os.environ.copy())
        asyproc.wait()
        asyrc = asyproc.returncode
        if asyrc != 0:
            print("****asy fail****")
        return asyrc

def getLabel(test):
    if test['function']=='gemm':
        return 'transA ' + test['transA']+ ' transB ' + test['transB']
    elif  test['function']=='axpy':
        return 'alpha '+str(test['alpha'])+' incx '+str(test['incx'])+' incy '+str(test['incy'])
    elif  test['function']=='gemv':
        return 'transA ' + test['transA']+' incx '+str(test['incx'])+' incy '+str(test['incy'])
    elif test['function'] in ['dot', 'copy', 'swap', 'ger', 'gerc', 'geru']:
        return 'incx '+str(test['incx'])+' incy '+str(test['incy'])
    elif test['function'] in ['asum', 'nrm2', 'scal']:
        return 'incx '+str(test['incx'])
    elif test['function']=='trsm':
        if test['side']=='R':
            return 'N/lda '+str(test['N'])+' alpha '+ str(test['alpha']) + ' side ' + str(test['side']) + ' uplo ' + str(test['uplo']) + ' transA ' + test['transA'] + ' diag ' + str(test['diag'])
        else:
            return 'M/lda/ldb '+str(test['M'])+' alpha '+ str(test['alpha']) + ' side ' + str(test['side']) + ' uplo ' + str(test['uplo']) + ' transA ' + test['transA'] + ' diag ' + str(test['diag'])
    else:
        print('Legend label not defined for '+test['function'])
        sys.exit(1)

def getXLabel(test):
    if test.function=='gemm':
        return 'M=N=K=lda=ldb=ldc'
    elif  test.function in ['axpy', 'asum', 'dot', 'copy', 'nrm2', 'scal', 'swap']:
        return 'N'
    elif  test.function in ['gemv', 'ger', 'gerc', 'geru']:
        return 'M=N=lda'
    elif  test.function=='trsm':
        if test.side == 'R':
            return 'M=ldb'
        else:
            return 'N'
    else:
        print('Xlabel not defined for ' + test.function)
        sys.exit(1)

def getFunctionPreFix(computeType):
    if "32_r" in computeType:
        return "s"
    elif "64_r" in computeType:
        return "d"
    elif "32_c" in computeType:
        return "c"
    elif "64_c" in computeType:
        return "z"
    elif "bf16_r" in computeType:
        return "bf"
    elif "f16_r" in computeType:
        return "h"
    else:
        print("Error - Cannot detect precision preFix: "+computeType)

def getDeviceSpecs(device, sclk):
    hwinfo = {}
    hwinfo["theoMaxCompute"] = -1
    hwinfo["sclk"] = int(sclk.split('M')[0])
    # print(hwinfo["sclk"])
    # print(hwinfo["sclk"]/1000.00 * 64 * 128)
    if 'Vega 20' in device:
        hwinfo["theoMaxCompute"] = hwinfo["sclk"]/1000.00 * 64 * 128 # 64 CU, 128 ops/ clk
        hwinfo["Bandwidth"] = 1000
        hwinfo["Device"] = 'Vega 20'
    elif 'Vega 10' in device:
        hwinfo["theoMaxCompute"] = hwinfo["sclk"]/1000.00 * 60 * 128
        hwinfo["Bandwidth"] = 484
        hwinfo["Device"] = 'Vega 10'
    else:
        print("Device type not supported or found - needed to display theoretical max")

    return hwinfo


def main(argv):
    dirA = "."
    dirB = None
    dryrun = False
    inputYaml = ""
    outdir = "."
    baseOutDir = "."
    speedup = False
    datatype = "gflops"
    # shortrun = False
    docformat = "pdf"
    devicenum = 0
    doAsy = False
    noFigures = False
    nsample = 10

    try:
        opts, args = getopt.getopt(argv,"hA:f:B:Tt:a:b:o:S:sg:d:N:i:X")
    except getopt.GetoptError:
        print("error in parsing arguments.")
        print(usage)
        sys.exit(2)
    for opt, arg in opts:
        if opt in ("-h"):
            print(usage)
            exit(0)
        elif opt in ("-A"):
            dirA = arg
        elif opt in ("-B"):
            dirB = arg
        elif opt in ("-i"):
            inputYaml = arg
        elif opt in ("-o"):
            outdir = arg
        elif opt in ("-b"):
            baseOutDir = arg
        elif opt in ("-T"):
            dryrun = True
        # elif opt in ("-s"):
        #     shortrun = True
        elif opt in ("-X"):
            noFigures = True
        elif opt in ("-g"):
            if int(arg) == 0:
                doAsy = False
            if int(arg) == 1:
                doAsy = True
        elif opt in ("-d"):
            devicenum = int(arg)
        elif opt in ("-N"):
            nsample = int(arg)
        elif opt in ("-S"):
            if int(arg) == 0:
                speedup = False
            if int(arg) == 1:
                speedup = True
        elif opt in ("-t"):
            if arg not in ["time", "gflops"]:
                print("data type must be time or gflops")
                print(usage)
                sys.exit(1)
            datatype = arg
        elif opt in ("-f"):
            goodvals = ["pdf", "docx"]
            if arg not in goodvals:
                print("error: format must in " + " ".join(goodvals))
                print(usage)
                sys.exit(1)
            docformat = arg

    if os.path.isfile(inputYaml)==False:
        print("unable to find input yaml file: " + inputYaml)
        sys.exit(1)

    print("dirA: "+ dirA)

    if not dryrun and not binaryisok(dirA, "rocblas-bench"):
        print("unable to find " + "rocblas-bench" + " in " + dirA)
        print("please specify with -A")
        sys.exit(1)

    dirlist = [[dirA, outdir]]
    if not dirB == None:
        print("dirB: "+ dirB)

        if not dryrun and not binaryisok(dirB, "rocblas-bench"):
            print("unable to find " + "rocblas-bench" + " in " + dirB)
            print("please specify with -B")
            sys.exit(1)

        if not os.path.exists(baseOutDir):
            os.makedirs(baseOutDir)

        dirlist.append([dirB, baseOutDir])

    elif dryrun:
        dirlist.append([dirB, baseOutDir])

    print("outdir: " + outdir)
    # if shortrun:
    #     print("short run")
    print("output format: " + docformat)
    print("device number: " + str(devicenum))

    if not os.path.exists(outdir):
        os.makedirs(outdir)

    rocBlasVersion = getVersion(dirA)
    sclk = getspecs.getsclk(devicenum)
    device = getspecs.getdeviceinfo(devicenum)

    if not dryrun:
        specs = "Host info:\n"
        specs += "\thostname: " + getspecs.gethostname() + "\n"
        specs += "\tcpu info: " + getspecs.getcpu() + "\n"
        specs += "\tram: " + getspecs.getram() + "\n"
        specs += "\tdistro: " + getspecs.getdistro() + "\n"
        specs += "\tkernel version: " + getspecs.getkernel() + "\n"
        specs += "\trocm version: " + getspecs.getrocmversion() + "\n"
        specs += "\t" + rocBlasVersion.version + "\n"
        specs += "Device info:\n"
        specs += "\tdevice: " + device + "\n"
        specs += "\tvbios version: " + getspecs.getvbios(devicenum) + "\n"
        specs += "\tvram: " + getspecs.getvram(devicenum) + "\n"
        specs += "\tperformance level: " + getspecs.getperflevel(devicenum) + "\n"
        specs += "\tsystem clock: " + sclk + "\n"
        specs += "\tmemory clock: " + getspecs.getmclk(devicenum) + "\n"

        with open(os.path.join(outdir, "specs.txt"), "w+") as f:
            f.write(specs)

    hwinfo = getDeviceSpecs(device, sclk)

    figs = []

    #load yaml then create fig for every test
    f = open(inputYaml, 'r')
    data = yamldata(f)
    f.close()

    #only generate data
    if noFigures:
        benchruns = []
        for tests in data.testcases:
            for test in tests:
                for idx, lwdir in enumerate(dirlist):
                    wdir = lwdir[0]
                    odir = lwdir[1]
                    label = getLabel(test)
                    benchruns.append( rundata(wdir, odir, idx, label, test, hwinfo) )
        for run in benchruns:
            print(" ".join(run.runcmd(nsample)))
            run.executerun(nsample)
        return

    #setup tests sorted by their respectice figures
    for tests in data.testcases:
        name = getFunctionPreFix(tests[0]['compute_type']) + tests[0]['function'].split('_')[0] + " Performance"
        fig = figure(name , name.replace('_', '\_'))
        for test in tests:
            for idx, lwdir in enumerate(dirlist):
                wdir = lwdir[0]
                odir = lwdir[1]
                label = getLabel(test)
                fig.runs.append( rundata(wdir, odir, idx, label,
                                            test, hwinfo) )
        figs.append(fig)

    #print and launch blas functions
    for fig in figs:
        print(fig.name)
        for run in fig.runs:
            if not dryrun:
                print(" ".join(run.runcmd(nsample)))
                run.executerun(nsample)

    #generate plots
    if doAsy:
        print("")
        for fig in figs:
            ncompare = len(dirlist) if speedup else 0
            print(fig.labels())
            print(fig.asycmd(outdir, docformat, datatype, ncompare))
            fig.executeasy(outdir, docformat, datatype, ncompare)

        if docformat == "pdf":
            maketex(figs, outdir, nsample)
        if docformat == "docx":
            makedocx(figs, outdir, nsample)

def binaryisok(dirname, progname):
    prog = os.path.join(dirname, progname)
    return os.path.isfile(prog)

def maketex(figs, outdir, nsample):

    header = '''\documentclass[12pt]{article}
\\usepackage{graphicx}
\\usepackage{url}
\\author{Wasiq Mahmood}
\\begin{document}
'''
    texstring = header

    texstring += "\n\\section{Introduction}\n"

    texstring += "Each data point represents the median of " + str(nsample) + " values, with error bars showing the 95\\% confidence interval for the median.\n\n"
    #TODO change message
    # texstring += "The following figures display the performance of the user specified rocBLAS functions \n\n"



    texstring += "\\vspace{1cm}\n"

    # texstring += "\\begin{tabular}{ll}"
    # texstring += labelA +" &\\url{"+ dirA+"} \\\\\n"
    # if not dirB == None:
    #     texstring += labelB +" &\\url{"+ dirB+"} \\\\\n"
    # texstring += "\\end{tabular}\n\n"

    # texstring += "\\vspace{1cm}\n"

    specfilename = os.path.join(outdir, "specs.txt")
    if os.path.isfile(specfilename):
        specs = ""
        with open(specfilename, "r") as f:
            specs = f.read()

        for line in specs.split("\n"):
            if line.startswith("Host info"):
                texstring += "\\noindent " + line
                texstring += "\\begin{itemize}\n"
            elif line.startswith("Device info"):
                texstring += "\\end{itemize}\n"
                texstring += line
                texstring += "\\begin{itemize}\n"
            else:
                if line.strip() != "":
                    texstring += "\\item " + line + "\n"
        texstring += "\\end{itemize}\n"
        texstring += "\n"

    texstring += "\\clearpage\n"

    texstring += "\n\\section{Figures}\n"

    for fig in figs:
        print(fig.filename(outdir, "pdf"))
        print(fig.caption)
        texstring += '''
\\centering
\\begin{figure}[htbp]
   \\includegraphics[width=\\textwidth]{'''
        texstring += fig.filename("", "pdf")
        texstring += '''}
   \\caption{''' + fig.caption + '''}
\\end{figure}
'''

    texstring += "\n\\end{document}\n"

    fname = os.path.join(outdir, 'figs.tex')

    with open(fname, 'w') as outfile:
        outfile.write(texstring)

    fout = open(os.path.join(outdir, "texcmd.log"), 'w+')
    ferr = open(os.path.join(outdir, "texcmd.err"), 'w+')

    latexcmd = ["latexmk", "-pdf", 'figs.tex']
    print(" ".join(latexcmd))
    texproc =  subprocess.Popen(latexcmd, cwd=outdir, stdout=fout, stderr=ferr,
                                env=os.environ.copy())
    texproc.wait()
    fout.close()
    ferr.close()
    texrc = texproc.returncode
    if texrc != 0:
        print("****tex fail****")

def pdf2emf(pdfname):
    svgname = pdfname.replace(".pdf",".svg")
    cmd_pdf2svg = ["pdf2svg", pdfname, svgname]
    proc = subprocess.Popen(cmd_pdf2svg, env=os.environ.copy())
    proc.wait()
    if proc.returncode != 0:
        print("pdf2svg failed!")
        sys.exit(1)

    emfname = pdfname.replace(".pdf",".emf")
    cmd_svg2emf = ["inkscape", svgname, "-M", emfname]
    proc = subprocess.Popen(cmd_svg2emf, env=os.environ.copy())
    proc.wait()
    if proc.returncode != 0:
        print("svg2emf failed!")
        sys.exit(1)

    return emfname

def makedocx(figs, outdir, nsample):
    import docx

    document = docx.Document()

    document.add_heading('rocBLAS benchmarks', 0)

    document.add_paragraph("Each data point represents the median of " + str(nsample) + " values, with error bars showing the 95% confidence interval for the median.")

    specfilename = os.path.join(outdir, "specs.txt")
    if os.path.isfile(specfilename):
        with open(specfilename, "r") as f:
            specs = f.read()
        for line in specs.split("\n"):
            document.add_paragraph(line)

    for fig in figs:
        print(fig.filename(outdir, "docx"))
        print(fig.caption)
        emfname = pdf2emf(fig.filename(outdir, "docx"))
        # NB: emf support does not work; adding the filename as a placeholder
        # document.add_paragraph(emfname)
        document.add_picture(emfname, width=docx.shared.Inches(6))
        document.add_paragraph((fig.caption).replace('\\', ''))

    document.save(os.path.join(outdir,'figs.docx'))

if __name__ == "__main__":
    main(sys.argv[1:])

